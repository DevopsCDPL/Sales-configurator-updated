"""
Forged-IDAS & CDPL Platform — Comprehensive Architecture Document Generator
Generates a professional Word document with diagrams (text-box based) and detailed architecture.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ─── Styles ────────────────────────────────────────────────────────────

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x0F, 0x1A, 0x2E)

# ─── Helper Functions ──────────────────────────────────────────────────

def add_colored_heading(text, level=1, color=RGBColor(0x0F, 0x1A, 0x2E)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0F1A2E"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_diagram_box(text, width=6.5):
    """Add a monospace text-based diagram in a shaded box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    # Use a table cell for shading background
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Inches(width)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    # Set border
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:color="CBD5E1"/>'
        f'<w:left w:val="single" w:sz="4" w:color="CBD5E1"/>'
        f'<w:bottom w:val="single" w:sz="4" w:color="CBD5E1"/>'
        f'<w:right w:val="single" w:sz="4" w:color="CBD5E1"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    for line in text.strip().split('\n'):
        para = cell.add_paragraph(line)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        for run in para.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    # Remove the first empty paragraph in the cell
    if cell.paragraphs[0].text == '':
        p_elem = cell.paragraphs[0]._element
        p_elem.getparent().remove(p_elem)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(f'📌 {text}')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

# ════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ════════════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FORGED-IDAS & CDPL PLATFORM')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x0F, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Comprehensive System Architecture Document')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1F, 0x7A, 0x63)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f'Version 1.0  •  {datetime.date.today().strftime("%B %d, %Y")}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta2.add_run('Cholan Dynamics Private Limited (CDPL)')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════

add_colored_heading('Table of Contents', 1)
toc_items = [
    '1. Executive Summary',
    '2. System Overview & High-Level Architecture',
    '3. Technology Stack',
    '4. Frontend Architecture',
    '   4.1 Application Structure',
    '   4.2 Routing & Navigation',
    '   4.3 State Management',
    '   4.4 Component Hierarchy',
    '   4.5 Platform Admin Module',
    '5. Backend Architecture',
    '   5.1 Server & Entry Point',
    '   5.2 Middleware Pipeline',
    '   5.3 API Route Map',
    '   5.4 Service Layer',
    '   5.5 Controller Layer',
    '6. Database Architecture',
    '   6.1 Entity Relationship Overview',
    '   6.2 Core Domain Models',
    '   6.3 Procurement Models',
    '   6.4 Enterprise & Security Models',
    '   6.5 Migration Strategy',
    '7. Authentication & Authorization',
    '   7.1 Authentication Flow',
    '   7.2 Role-Based Access Control (RBAC)',
    '   7.3 Multi-Tenant Isolation',
    '8. API Architecture',
    '   8.1 REST API Design',
    '   8.2 API Security',
    '9. Data Flow Diagrams',
    '   9.1 User Login Flow',
    '   9.2 Project Lifecycle Flow',
    '   9.3 Procurement Flow',
    '   9.4 Multi-Tenant Data Flow',
    '10. Deployment Architecture',
    '   10.1 Infrastructure Diagram',
    '   10.2 CI/CD Pipeline',
    '   10.3 Environment Configuration',
    '11. Security Architecture',
    '12. Module Breakdown',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════

add_colored_heading('1. Executive Summary', 1)
doc.add_paragraph(
    'Forged-IDAS (Integrated Digital Administrative System) is a full-stack, multi-tenant '
    'SaaS platform built for manufacturing companies to manage their end-to-end business '
    'operations — from client acquisition and project estimation through procurement, '
    'production, quality control, logistics, and invoicing.'
)
doc.add_paragraph(
    'The CDPL Platform is the super-admin layer that sits above Forged-IDAS, providing '
    'multi-tenant company management, subscription control, platform-wide analytics, '
    'and centralized user administration for the SaaS operator (Cholan Dynamics Private Limited).'
)
doc.add_paragraph(
    'Together, the two layers form a complete enterprise platform:'
)
add_bullet('Forged-IDAS', bold_prefix='Forged-IDAS: ')
p = doc.paragraphs[-1]
p.clear()
run = p.add_run('Forged-IDAS: ')
run.bold = True
p.add_run('Tenant-level ERP for manufacturing operations')

add_bullet('', bold_prefix='CDPL Platform: ')
p = doc.paragraphs[-1]
p.clear()
run = p.add_run('CDPL Platform: ')
run.bold = True
p.add_run('Platform-level administration for multi-tenant SaaS management')

# ════════════════════════════════════════════════════════════════════════
#  2. SYSTEM OVERVIEW & HIGH-LEVEL ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_colored_heading('2. System Overview & High-Level Architecture', 1)
doc.add_paragraph('The system follows a classic three-tier web architecture with clear separation of concerns:')

add_colored_heading('2.1 High-Level Architecture Diagram', 2)
add_diagram_box('''
┌─────────────────────────────────────────────────────────────────────────────┐
│                              USERS / BROWSERS                               │
│         Platform Admin  │  Super Admin  │  Admin  │  User  │  Sales Eng     │
└────────────────────┬────────────────────────────────────────────────────────┘
                     │  HTTPS
                     ▼
┌─────────────────────────────────────────────────────────────────────────────┐
│                         RAILWAY CLOUD PLATFORM                              │
│                                                                             │
│  ┌──────────────────────────────┐   ┌──────────────────────────────────┐    │
│  │     FRONTEND SERVICE         │   │       BACKEND SERVICE            │    │
│  │  ┌────────────────────────┐  │   │  ┌──────────────────────────┐   │    │
│  │  │   React 18 + TypeScript│  │   │  │   Express.js + Node 20   │   │    │
│  │  │   Material-UI 5.15     │  │   │  │   Sequelize 6 ORM        │   │    │
│  │  │   React Router 6       │◄─┼───┼─►│   JWT Authentication     │   │    │
│  │  │   Axios HTTP Client    │  │   │  │   Multi-tenant Middleware │   │    │
│  │  │   Lucide Icons         │  │   │  │   PDFKit / ExcelJS       │   │    │
│  │  └────────────────────────┘  │   │  └──────────┬───────────────┘   │    │
│  │  Served by: node start.js    │   │             │                    │    │
│  │  (Runtime API URL injection) │   │             │ Sequelize          │    │
│  └──────────────────────────────┘   │             ▼                    │    │
│                                      │  ┌──────────────────────────┐   │    │
│                                      │  │   PostgreSQL Database    │   │    │
│                                      │  │   (Railway Managed)      │   │    │
│                                      │  │   ─────────────────────  │   │    │
│                                      │  │   59 Models / 61 Migr.   │   │    │
│                                      │  │   Multi-tenant Scoped    │   │    │
│                                      │  └──────────────────────────┘   │    │
│                                      └──────────────────────────────────┘    │
└─────────────────────────────────────────────────────────────────────────────┘
''')

add_colored_heading('2.2 Dual-Application Architecture', 2)
add_diagram_box('''
┌─────────────────────────────────────────────────────────────────┐
│                    SINGLE CODEBASE (Monorepo)                   │
│                                                                 │
│  ┌─────────────────────┐     ┌─────────────────────────────┐   │
│  │   CDPL PLATFORM      │     │     FORGED-IDAS (ERP)       │   │
│  │   /platform-admin/*  │     │     /* (all other routes)   │   │
│  │                      │     │                             │   │
│  │  • Company Mgmt      │     │  • Projects & Estimates     │   │
│  │  • User Mgmt         │     │  • Sales Orders             │   │
│  │  • Subscriptions     │     │  • Work Orders              │   │
│  │  • Billing           │     │  • Procurement              │   │
│  │  • Platform Analytics│     │  • Vendors & Clients        │   │
│  │  • Audit Logs        │     │  • Quality & Logistics      │   │
│  │  • API Keys          │     │  • Invoicing                │   │
│  │  • Integrations      │     │  • File Manager             │   │
│  │                      │     │  • Analytics & Reports      │   │
│  │  Guard:              │     │  • Access Control (RBAC)    │   │
│  │  PlatformAdminGuard  │     │  • Enterprise Features      │   │
│  │  (platform_admin     │     │                             │   │
│  │   role only)         │     │  Guards: ProtectedRoute,    │   │
│  │                      │     │  RBACRoute, CoAdminRoute    │   │
│  └─────────────────────┘     └─────────────────────────────┘   │
└─────────────────────────────────────────────────────────────────┘
''')

# ════════════════════════════════════════════════════════════════════════
#  3. TECHNOLOGY STACK
# ════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_colored_heading('3. Technology Stack', 1)

add_colored_heading('3.1 Frontend', 2)
add_table(
    ['Technology', 'Version', 'Purpose'],
    [
        ['React', '18.2', 'UI library (SPA)'],
        ['TypeScript', '5.9', 'Type safety'],
        ['Material-UI (MUI)', '5.15', 'Component library'],
        ['React Router', '6.21', 'Client-side routing'],
        ['Axios', '1.6', 'HTTP client'],
        ['Recharts', '3.7', 'Charts & data visualization'],
        ['Lucide React', '0.577', 'Icon system'],
        ['Framer Motion', '12.35', 'Animations'],
        ['React Hook Form', '7.49', 'Form management'],
        ['Day.js', '-', 'Date formatting'],
    ],
    [2.5, 1, 3]
)

doc.add_paragraph()
add_colored_heading('3.2 Backend', 2)
add_table(
    ['Technology', 'Version', 'Purpose'],
    [
        ['Node.js', '20 LTS', 'Runtime environment'],
        ['Express.js', '4.18.2', 'Web framework'],
        ['Sequelize', '6.35.2', 'ORM (PostgreSQL)'],
        ['PostgreSQL (pg)', '8.11.3', 'Database driver'],
        ['JSON Web Token', '9.0.2', 'Authentication'],
        ['bcryptjs', '2.4.3', 'Password hashing'],
        ['Multer', '1.4.5', 'File upload handling'],
        ['PDFKit', '0.14.0', 'PDF generation'],
        ['ExcelJS', '4.4.0', 'Excel export'],
        ['Nodemailer', '8.0.2', 'Email service'],
        ['Helmet', '8.1.0', 'HTTP security headers'],
        ['express-rate-limit', '8.2.1', 'API rate limiting'],
    ],
    [2.5, 1, 3]
)

doc.add_paragraph()
add_colored_heading('3.3 Infrastructure', 2)
add_table(
    ['Component', 'Provider', 'Details'],
    [
        ['Hosting', 'Railway', 'PaaS with auto-deploy from Git'],
        ['Database', 'Railway PostgreSQL', 'Managed PostgreSQL instance'],
        ['Build System', 'Nixpacks', 'Auto-detected build pipeline'],
        ['Version Control', 'GitHub', 'DevopsCDPL/forged-idas'],
        ['Branch Strategy', 'dev → main', 'dev auto-deploys, main protected'],
        ['Frontend Runtime', 'start.js', 'Dynamic API URL injection'],
    ],
    [2, 2, 2.5]
)

# ════════════════════════════════════════════════════════════════════════
#  4. FRONTEND ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_colored_heading('4. Frontend Architecture', 1)

add_colored_heading('4.1 Application Structure', 2)
add_diagram_box('''
frontend/src/
├── App.tsx                    # Root component + route definitions
├── theme.ts                   # MUI theme configuration
├── index.tsx                  # React DOM entry point
│
├── config/
│   ├── rolePermissions.ts     # RBAC rules, path access control
│   └── platformAdminMenu.ts   # Platform admin sidebar config
│
├── contexts/
│   ├── AuthContext.tsx         # Authentication state (user, token)
│   ├── ThemeContext.tsx        # Theme toggle (dark/light)
│   └── NotificationContext.tsx # Toast notification system
│
├── services/
│   ├── api.ts                 # Axios instance (base URL, JWT, 401 handling)
│   ├── authService.ts         # Login, register, profile, password reset
│   ├── projectService.ts      # Project CRUD operations
│   ├── clientService.ts       # Client management
│   ├── vendorService.ts       # Vendor management
│   ├── estimateService.ts     # Estimate operations
│   ├── procurementService.ts  # Procurement workflows
│   ├── partService.ts         # Parts master
│   ├── materialService.ts     # Material management
│   ├── stockService.ts        # Inventory operations
│   └── ... (15+ service modules)
│
├── types/
│   └── index.ts               # TypeScript interfaces (User, Project, etc.)
│
├── utils/
│   ├── calculations.ts        # Math helpers
│   └── documentUtils.ts       # Document generation helpers
│
├── components/
│   ├── Layout/
│   │   └── Layout.tsx         # Main app shell (sidebar + topbar + outlet)
│   ├── PlatformAdmin/
│   │   ├── PlatformAdminLayout.tsx    # Platform admin shell
│   │   └── PlatformAdminSidebar.tsx   # Enterprise sidebar component
│   ├── ProjectTabs/           # 12 tab components for project detail
│   ├── ErrorBoundary.tsx
│   ├── GlobalSearch.tsx
│   ├── CommandPalette.tsx
│   ├── AIAssistant.tsx
│   └── ... (30+ components)
│
└── pages/
    ├── LoginPage.tsx, DashboardPage.tsx
    ├── ProjectsPage.tsx, ProjectDetailPage.tsx
    ├── ClientsPage.tsx, VendorsPage.tsx
    ├── PartsMasterPage.tsx, RawMaterialMasterPage.tsx
    ├── AccessControlPage.tsx, SettingsPage.tsx
    ├── platform-admin/
    │   ├── PlatformDashboardPage.tsx
    │   ├── PlatformCompaniesPage.tsx
    │   ├── PlatformUsersPage.tsx
    │   └── PlaceholderPage.tsx (12 placeholders)
    └── ... (35+ pages)
''')

add_colored_heading('4.2 Routing & Navigation', 2)
doc.add_paragraph('The application uses React Router v6 with nested route layouts and multiple route guards:')

add_diagram_box('''
Route Tree:
═══════════════════════════════════════════════════════════════
/login                          → LoginPage (public)
/forgot-password                → ForgotPasswordPage (public)
│
/ (Layout)                      → ProtectedRoute guard
├── /                           → DashboardPage
├── /projects                   → ProjectsPage
├── /projects/:id               → ProjectDetailPage
├── /clients                    → ClientsPage        [RBACRoute]
├── /vendors                    → VendorsPage         [RBACRoute]
├── /vendor-po                  → VendorPOPage        [RBACRoute]
├── /parts-master               → PartsMasterPage     [RBACRoute]
├── /raw-materials              → RawMaterialMasterPage[RBACRoute]
├── /procurement                → MgmtProcurementPage [RBACRoute]
├── /material-stock             → MaterialStockPage   [RBACRoute]
├── /file-manager               → FileManagerPage
├── /access-control/*           → AccessControlPage   [CoAdminRoute]
├── /settings                   → SettingsPage
├── /recycle-bin                → RecycleBinPage      [CoAdminRoute]
├── /sessions                   → SessionMonitoring   [RBACRoute]
├── /custom-roles               → CustomRoleBuilder   [RBACRoute]
├── /analytics                  → BusinessAnalyticsPage[RBACRoute]
└── /messages                   → ChatPage
│
/platform-admin (PlatformAdminLayout)  → PlatformAdminGuard
├── /                           → PlatformDashboardPage
├── /companies                  → PlatformCompaniesPage
├── /users                      → PlatformUsersPage
├── /admins                     → PlaceholderPage
├── /teams                      → PlaceholderPage
├── /activity-logs              → PlaceholderPage
├── /audit-logs                 → PlaceholderPage
├── /notifications              → PlaceholderPage
├── /reports                    → PlaceholderPage
├── /insights                   → PlaceholderPage
├── /subscriptions              → PlaceholderPage
├── /billing                    → PlaceholderPage
├── /integrations               → PlaceholderPage
├── /api-keys                   → PlaceholderPage
├── /settings                   → PlatformSettingsPage
└── /recycle-bin                → PlaceholderPage
''')

add_colored_heading('4.3 State Management', 2)
doc.add_paragraph('The application uses React Context API for global state (no Redux):')

add_table(
    ['Context', 'State', 'Purpose'],
    [
        ['AuthContext', 'user, token, isAuthenticated, isLoading', 'Authentication state with JWT management'],
        ['ThemeContext', 'theme mode', 'Dark/light theme toggle'],
        ['NotificationContext', 'toast queue', 'Global notification system'],
    ],
    [2, 2.5, 2]
)

add_colored_heading('4.4 Component Hierarchy', 2)
add_diagram_box('''
<App>
├── <ErrorBoundary>
│   └── <Routes>
│       ├── <LoginPage />
│       │
│       ├── <ProtectedRoute>
│       │   └── <Layout>                        ← Main App Shell
│       │       ├── Sidebar Navigation
│       │       ├── Top AppBar (search, user menu)
│       │       ├── <CommandPalette />
│       │       ├── <GlobalSearch />
│       │       └── <Outlet />                   ← Page Content
│       │           ├── <DashboardPage>
│       │           │   ├── <StatCards />
│       │           │   ├── <ProjectChart />
│       │           │   └── <RecentProjects />
│       │           ├── <ProjectDetailPage>
│       │           │   └── <ProjectTabs>
│       │           │       ├── InfoTab
│       │           │       ├── EstimationTab
│       │           │       ├── QuotationTab
│       │           │       ├── SalesOrderTab
│       │           │       ├── WorkOrderTab
│       │           │       ├── VendorPOTab
│       │           │       ├── ProductionTab
│       │           │       ├── QualityTab
│       │           │       ├── LogisticsTab
│       │           │       ├── InvoiceTab
│       │           │       ├── DocumentsTab
│       │           │       └── AnalyticsTab
│       │           └── ...other pages
│       │
│       └── <PlatformAdminGuard>
│           └── <PlatformAdminLayout>            ← Platform Admin Shell
│               ├── <PlatformAdminSidebar />
│               ├── TopBar
│               └── <Outlet />
│                   ├── <PlatformDashboardPage />
│                   ├── <PlatformCompaniesPage />
│                   └── ...
''')

add_colored_heading('4.5 Platform Admin Sidebar Architecture', 2)
doc.add_paragraph('The Platform Admin sidebar uses a config-driven, role-based approach:')

add_diagram_box('''
┌──────────────────────────────────────────────────────────┐
│  platformAdminMenu.ts (Config)                           │
│  ┌────────────────────────────────────────────────────┐  │
│  │ PLATFORM_MENU: PlatformMenuSection[]               │  │
│  │  ├─ MAIN:       Dashboard                          │  │
│  │  ├─ MANAGEMENT: Companies, Users, Admins, Teams    │  │
│  │  ├─ OPERATIONS: Activity Logs, Audit Logs, Notif.  │  │
│  │  ├─ ANALYTICS:  Reports, Insights                  │  │
│  │  ├─ SYSTEM:     Subscriptions, Billing, API Keys   │  │
│  │  └─ SAFETY:     Recycle Bin                        │  │
│  └────────────────────────────────────────────────────┘  │
│                        │                                  │
│                        ▼                                  │
│  filterMenuByRole(role) ─── Filters items by PlatformRole │
│                        │                                  │
│                        ▼                                  │
│  ┌─────────────────────────────────────────────────────┐ │
│  │  PlatformAdminSidebar.tsx                           │ │
│  │  ├─ Header (logo + collapse toggle)                 │ │
│  │  ├─ SidebarSection[] (collapsible groups)           │ │
│  │  │   └─ SidebarItem[] (icon + label + badge)        │ │
│  │  └─ Footer (avatar + logout + version)              │ │
│  └─────────────────────────────────────────────────────┘ │
└──────────────────────────────────────────────────────────┘

Role Visibility:
  Owner ──► All 16 items
  Admin ──► 14 items (no Billing, API Keys)
  User  ──► 3 items  (Dashboard, Companies, Users)
''')

# ════════════════════════════════════════════════════════════════════════
#  5. BACKEND ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_colored_heading('5. Backend Architecture', 1)

add_colored_heading('5.1 Server Architecture', 2)
add_diagram_box('''
backend/src/
├── index.js                 # Express server entry point
├── config/
│   └── database.js          # Sequelize config (PostgreSQL, SSL, pooling)
│
├── middleware/
│   ├── auth.js              # JWT verification, req.user attachment
│   ├── tenant.js            # Multi-tenant scoping & isolation
│   ├── tenantScope.js       # Automatic query filtering by company_id
│   └── validate.js          # Express-validator integration
│
├── models/                  # 59 Sequelize models
│   ├── User.js, Company.js, Client.js, Vendor.js
│   ├── Project.js, Estimate.js, SalesOrder.js, WorkOrder.js
│   ├── Material.js, RawMaterial.js, Part.js, Stock.js
│   ├── ProcurementRFQ.js, ProcurementPO.js
│   ├── Document.js, AuditLog.js, Session.js
│   └── ... (59 total)
│
├── routes/                  # 44 route modules
│   ├── authRoutes.js, userRoutes.js
│   ├── projectRoutes.js, estimateRoutes.js
│   ├── vendorRoutes.js, clientRoutes.js
│   ├── procurementRoutes.js, materialRoutes.js
│   └── ... (44 total, all mounted under /api)
│
├── controllers/             # 40+ request handlers
│   └── (Thin layer → delegates to services)
│
├── services/                # 39+ business logic services
│   ├── authService.js, projectService.js
│   ├── estimateService.js, procurementService.js
│   ├── documentNumberingService.js
│   ├── excelExportService.js
│   └── documentDataMapper/  # PDF data transformation
│       ├── cocMapper.js, invoiceMapper.js
│       ├── rfqMapper.js, productionMapper.js
│       └── vendorPOMapper.js
│
├── scripts/                 # Database maintenance
└── data/                    # Static reference data
''')

add_colored_heading('5.2 Middleware Pipeline', 2)
doc.add_paragraph('Every API request passes through the following middleware chain:')

add_diagram_box('''
Request Flow:
═══════════════════════════════════════════════════════════════

  HTTP Request
       │
       ▼
  ┌─────────────┐
  │   CORS      │  Allow Railway, localhost, custom origins
  └──────┬──────┘
         ▼
  ┌─────────────┐
  │   Helmet    │  Security headers (CSP, HSTS, X-Frame-Options)
  └──────┬──────┘
         ▼
  ┌─────────────┐
  │ Rate Limit  │  express-rate-limit (per IP)
  └──────┬──────┘
         ▼
  ┌─────────────┐
  │ Body Parser │  JSON (10mb limit) + URL-encoded
  └──────┬──────┘
         ▼
  ┌─────────────┐
  │   /health   │──► 200 OK (bypass auth)
  └──────┬──────┘
         ▼
  ┌─────────────┐
  │  auth.js    │  JWT verify → req.user { id, email, role, company_id }
  └──────┬──────┘
         ▼
  ┌─────────────┐
  │ tenant.js   │  Scope to company_id, prevent tenant spoofing
  └──────┬──────┘
         ▼
  ┌─────────────┐
  │ validate.js │  Express-validator rules (body, query, params)
  └──────┬──────┘
         ▼
  ┌─────────────┐
  │  Route      │──► Controller ──► Service ──► Sequelize ──► PostgreSQL
  └─────────────┘
         │
         ▼
  JSON Response
''')

add_colored_heading('5.3 API Route Map', 2)

add_table(
    ['Route Prefix', 'Module', 'Auth', 'Role Restriction'],
    [
        ['/api/auth', 'Authentication', 'Public', 'None'],
        ['/api/users', 'User Management', 'JWT', 'Admin+'],
        ['/api/companies', 'Company Management', 'JWT', 'Main Admin+'],
        ['/api/clients', 'Client Management', 'JWT', 'RBAC'],
        ['/api/vendors', 'Vendor Management', 'JWT', 'RBAC'],
        ['/api/projects', 'Project Management', 'JWT', 'All roles'],
        ['/api/estimates', 'Cost Estimation', 'JWT', 'All roles'],
        ['/api/sales-orders', 'Sales Orders', 'JWT', 'All roles'],
        ['/api/work-orders', 'Work Orders', 'JWT', 'All roles'],
        ['/api/quality', 'Quality Records', 'JWT', 'All roles'],
        ['/api/logistics', 'Shipping/Delivery', 'JWT', 'All roles'],
        ['/api/documents', 'Document Mgmt', 'JWT', 'All roles'],
        ['/api/materials', 'Material Master', 'JWT', 'RBAC'],
        ['/api/parts', 'Parts Master', 'JWT', 'RBAC'],
        ['/api/raw-materials', 'Raw Materials', 'JWT', 'RBAC'],
        ['/api/material-stock', 'Inventory', 'JWT', 'RBAC'],
        ['/api/procurement', 'Procurement RFQ/PO', 'JWT', 'RBAC'],
        ['/api/mgmt-procurement', 'Internal Procurement', 'JWT', 'Admin+'],
        ['/api/vendor-po', 'Vendor POs', 'JWT', 'RBAC'],
        ['/api/vendor-procurement', 'Vendor RFQs', 'JWT', 'RBAC'],
        ['/api/permissions', 'Permission Mgmt', 'JWT', 'Main Admin'],
        ['/api/custom-roles', 'Custom Roles', 'JWT', 'Owner/Co-Owner'],
        ['/api/audit-logs', 'Audit Trail', 'JWT', 'Admin+'],
        ['/api/sessions', 'Session Monitor', 'JWT', 'Owner/Co-Owner'],
        ['/api/approvals', 'Workflow Approvals', 'JWT', 'Owner/Co-Owner'],
        ['/api/analytics', 'Analytics', 'JWT', 'Admin+'],
        ['/api/settings', 'System Settings', 'JWT', 'Admin+'],
        ['/api/recycle-bin', 'Soft Deletes', 'JWT', 'Main Admin'],
        ['/api/platform-admin', 'Platform Admin', 'JWT', 'Platform Admin only'],
        ['/api/ai-assistant', 'AI Features', 'JWT', 'All roles'],
        ['/api/chat', 'Messaging', 'JWT', 'All roles'],
        ['/api/search', 'Global Search', 'JWT', 'All roles'],
    ],
    [1.8, 1.5, 0.8, 1.8]
)

add_colored_heading('5.4 Service Layer Pattern', 2)
doc.add_paragraph('The backend follows a strict Controller → Service → Model pattern:')

add_diagram_box('''
┌──────────────┐     ┌──────────────┐     ┌──────────────┐     ┌──────────┐
│   Route      │────►│  Controller  │────►│   Service    │────►│  Model   │
│ (validation) │     │ (req/res)    │     │ (business    │     │(Sequelize│
│              │     │              │     │  logic)      │     │  query)  │
└──────────────┘     └──────────────┘     └──────────────┘     └──────────┘
                                                │
                                    ┌───────────┼───────────┐
                                    ▼           ▼           ▼
                              PDF Gen     Excel Export   Email
                             (PDFKit)     (ExcelJS)    (Nodemailer)
''')

# ════════════════════════════════════════════════════════════════════════
#  6. DATABASE ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_colored_heading('6. Database Architecture', 1)

add_colored_heading('6.1 Entity Relationship Overview', 2)
add_diagram_box('''
                              ENTITY RELATIONSHIP DIAGRAM
═══════════════════════════════════════════════════════════════════════════

                            ┌─────────────┐
                            │   COMPANY   │
                            │─────────────│
                            │ id (PK)     │
                            │ company_name│
                            │ company_code│
                            │ plan        │
                            │ is_active   │
                            │ subscription│
                            └──────┬──────┘
                    ┌──────────────┼──────────────────────┐
                    │              │                      │
                    ▼              ▼                      ▼
             ┌───────────┐  ┌───────────┐         ┌───────────┐
             │   USER    │  │  CLIENT   │         │  VENDOR   │
             │───────────│  │───────────│         │───────────│
             │ id (PK)   │  │ id (PK)   │         │ id (PK)   │
             │ name      │  │ name      │         │ name      │
             │ email     │  │ company   │         │ category  │
             │ role      │  │ contact   │         │ contact   │
             │ company_id│  │ company_id│         │ company_id│
             └─────┬─────┘  └─────┬─────┘         └─────┬─────┘
                   │              │                      │
                   │              ▼                      │
                   │       ┌─────────────┐               │
                   │       │   PROJECT   │               │
                   │       │─────────────│               │
                   │       │ id (PK)     │               │
                   ├──────►│ client_id   │               │
                   │       │ status      │               │
                   │       │ company_id  │               │
                   │       └──────┬──────┘               │
                   │    ┌────┬───┬┴──┬────┬────┐         │
                   │    ▼    ▼   ▼   ▼    ▼    ▼         │
                   │  EST  S.O. W.O. QR  DOC  INV        │
                   │                                     │
                   │         ┌──────────────┐             │
                   │         │   MATERIAL   │◄────────────┘
                   │         │──────────────│  (vendor_id)
                   │         │ PART         │
                   │         │ RAW MATERIAL │
                   │         │ STOCK        │
                   │         └──────┬───────┘
                   │                │
                   │                ▼
                   │      ┌──────────────────┐
                   │      │  PROCUREMENT     │
                   │      │  ├─ RFQ          │
                   │      │  ├─ RFQ Items    │
                   │      │  ├─ Vendor Quotes│
                   │      │  ├─ PO           │
                   │      │  └─ PO Items     │
                   │      └──────────────────┘
                   │
                   ▼
           ┌──────────────────┐
           │  ENTERPRISE      │
           │  ├─ Permission   │
           │  ├─ CustomRole   │
           │  ├─ AuditLog     │
           │  ├─ Session      │
           │  ├─ Approval     │
           │  └─ RiskScore    │
           └──────────────────┘

Legend: EST=Estimate, S.O.=SalesOrder, W.O.=WorkOrder,
        QR=QualityRecord, DOC=Document, INV=Invoice
''')

add_colored_heading('6.2 Core Domain Models (59 Total)', 2)
add_table(
    ['Category', 'Models', 'Count'],
    [
        ['Core Business', 'User, Company, Client, Vendor, Project', '5'],
        ['Project Workflow', 'Estimate, SalesOrder, WorkOrder, QualityRecord, Invoice', '5'],
        ['Documents', 'Document, FileManagerFolder', '2'],
        ['Materials', 'Material, RawMaterial, Part, PartDimension, PartTemplate', '5'],
        ['Inventory', 'MaterialStock, MaterialTransaction, Stock', '3'],
        ['Procurement', 'ProcurementRFQ, RFQItem, RFQVendor, VendorQuote, ProcurementPO, POItem', '6'],
        ['Vendor Workflows', 'VendorPurchaseOrder, VendorPOItem, VendorRFQ, VendorPO, VendorMaterial', '5'],
        ['Mgmt Procurement', 'MgmtProcurementRFQ, MgmtProcurementPO', '2'],
        ['Enterprise/RBAC', 'Permission, PermissionTemplate, CustomRole, ApprovalWorkflow', '4'],
        ['Security/Audit', 'AuditLog, Session, RiskScore, LoginHistory, OtpToken', '5'],
        ['Communication', 'Conversation, ConversationParticipant, Message', '3'],
        ['Analytics', 'ProjectAnalytics, ActivityTimeline', '2'],
        ['Integration', 'ApiToken, Webhook, Setting', '3'],
        ['Bundles', 'RFQBundle, RFQBundleItem', '2'],
    ],
    [1.8, 3.5, 0.7]
)

add_colored_heading('6.3 Project Status State Machine', 2)
add_diagram_box('''
Project Lifecycle (status field):
═══════════════════════════════════════════════════════════════

  draft ──► estimated ──► quoted ──► order_confirmed
                                         │
                                         ▼
                                    in_production
                                         │
                                         ▼
                                     inspected
                                         │
                                         ▼
                                      shipped
                                         │
                                         ▼
                                      closed

  (Any state) ──► cancelled (soft delete available)
''')

add_colored_heading('6.4 Migration Strategy', 2)
doc.add_paragraph('The database uses Sequelize CLI migrations (61 migration files) with the naming convention:')
doc.add_paragraph('YYYYMMDDHHMMSS-description.js')
doc.add_paragraph('Migrations are idempotent (check-before-create pattern) and support both up() and down() operations.')

# ════════════════════════════════════════════════════════════════════════
#  7. AUTHENTICATION & AUTHORIZATION
# ════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_colored_heading('7. Authentication & Authorization', 1)

add_colored_heading('7.1 Authentication Flow', 2)
add_diagram_box('''
Login Flow:
═══════════════════════════════════════════════════════════════

  Browser                    Backend                     Database
    │                          │                            │
    │  POST /api/auth/login    │                            │
    │  {email, password}       │                            │
    │─────────────────────────►│                            │
    │                          │  Find user by email        │
    │                          │───────────────────────────►│
    │                          │  ◄─── User record          │
    │                          │                            │
    │                          │  bcrypt.compare(password)  │
    │                          │  ──── Verify ────          │
    │                          │                            │
    │                          │  Check: is_active?         │
    │                          │  Check: account_locked?    │
    │                          │  Check: deleted_at null?   │
    │                          │                            │
    │                          │  jwt.sign({userId}, secret)│
    │                          │  ──── Generate Token ────  │
    │                          │                            │
    │                          │  Log to LoginHistory       │
    │                          │───────────────────────────►│
    │                          │                            │
    │  ◄─── { user, token }    │                            │
    │                          │                            │
    │  Store token in          │                            │
    │  localStorage            │                            │
    │                          │                            │
    │  Subsequent requests:    │                            │
    │  Authorization: Bearer   │                            │
    │  <token>                 │                            │
    │─────────────────────────►│                            │
    │                          │  auth.js middleware         │
    │                          │  jwt.verify(token)         │
    │                          │  Attach req.user           │
    │                          │  {id,email,role,company_id}│
    │                          │                            │
''')

add_colored_heading('7.2 Role-Based Access Control (RBAC)', 2)

add_table(
    ['Role', 'Hierarchy', 'Access Scope', 'Can Create'],
    [
        ['platform_admin', 'Super (all companies)', 'All companies, platform admin panel', 'All roles'],
        ['main_admin', 'Level 100', 'Full company access, RBAC mgmt, Enterprise features', 'All except platform_admin'],
        ['admin', 'Level 70', 'Company data, master data, no Admin/Enterprise', 'user, sales_engineer'],
        ['sales_engineer', 'Level 30', 'Projects, Clients, Messages, File Manager', 'None'],
        ['user', 'Level 10', 'Projects, limited modules', 'None'],
    ],
    [1.3, 1.2, 2.5, 1.5]
)

doc.add_paragraph()
doc.add_paragraph('Co-Admin System: Within main_admin role, users can be assigned as Owner, Co-Owner, or Backup via the Co-Admin assignment system stored in the Settings table.')

add_colored_heading('7.3 Multi-Tenant Isolation', 2)
add_diagram_box('''
Multi-Tenant Architecture:
═══════════════════════════════════════════════════════════════

  ┌────────────────────────────────────────────────────────┐
  │                   SHARED DATABASE                       │
  │                                                        │
  │  ┌────────────┐  ┌────────────┐  ┌────────────┐       │
  │  │ Company A  │  │ Company B  │  │ Company C  │       │
  │  │ (Tenant 1) │  │ (Tenant 2) │  │ (Tenant 3) │       │
  │  │            │  │            │  │            │       │
  │  │ Users      │  │ Users      │  │ Users      │       │
  │  │ Projects   │  │ Projects   │  │ Projects   │       │
  │  │ Vendors    │  │ Vendors    │  │ Vendors    │       │
  │  │ ...        │  │ ...        │  │ ...        │       │
  │  └────────────┘  └────────────┘  └────────────┘       │
  │                                                        │
  │  Isolation Method: Row-level filtering via company_id  │
  │  Middleware: tenant.js + tenantScope.js                 │
  │  Platform Admin: Bypasses tenant scope (cross-company)  │
  └────────────────────────────────────────────────────────┘

  Anti-Spoofing:
  • Request body company_id overridden by JWT company_id
  • Query params filtered by middleware
  • Platform admin explicitly opted in (no auto-scope)
''')

# ════════════════════════════════════════════════════════════════════════
#  8. API ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_colored_heading('8. API Architecture', 1)

add_colored_heading('8.1 REST API Design', 2)
doc.add_paragraph('All APIs follow RESTful conventions:')
add_bullet('Base URL: https://forged-idas-production.up.railway.app/api')
add_bullet('Format: JSON request/response bodies')
add_bullet('Authentication: Bearer token in Authorization header')
add_bullet('Pagination: ?page=1&limit=20 (offset-based)')
add_bullet('Error format: { error: string, details?: any }')

add_colored_heading('8.2 API Security Measures', 2)
add_table(
    ['Security Layer', 'Implementation', 'Details'],
    [
        ['HTTPS', 'Railway TLS', 'All traffic encrypted in transit'],
        ['CORS', 'cors middleware', 'Whitelist of allowed origins'],
        ['Helmet', 'helmet middleware', 'CSP, HSTS, X-Frame-Options, etc.'],
        ['Rate Limiting', 'express-rate-limit', 'Per-IP request throttling'],
        ['JWT Auth', 'jsonwebtoken', 'Signed tokens with expiry'],
        ['Password Hash', 'bcryptjs (10 rounds)', 'Argon-style hashing'],
        ['Input Validation', 'express-validator', 'Whitelist validation on all inputs'],
        ['Tenant Isolation', 'tenant.js middleware', 'Prevents cross-tenant data access'],
        ['File Upload', 'multer + path validation', 'Restricted file types and sizes'],
        ['Account Lock', 'Failed attempt tracking', 'Auto-lock after failed logins'],
    ],
    [1.5, 1.8, 3]
)

# ════════════════════════════════════════════════════════════════════════
#  9. DATA FLOW DIAGRAMS
# ════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_colored_heading('9. Data Flow Diagrams', 1)

add_colored_heading('9.1 Project Lifecycle Flow', 2)
add_diagram_box('''
Complete Project Flow:
═══════════════════════════════════════════════════════════════

  Client Request
       │
       ▼
  ┌──────────────┐    ┌──────────────┐    ┌──────────────┐
  │   CREATE     │───►│  ESTIMATION  │───►│  QUOTATION   │
  │   PROJECT    │    │  (Revisions) │    │  (PDF Gen)   │
  │   (draft)    │    │  (estimated) │    │  (quoted)    │
  └──────────────┘    └──────────────┘    └──────────────┘
                                                │
                               Client Approves  │
                                                ▼
  ┌──────────────┐    ┌──────────────┐    ┌──────────────┐
  │   INVOICE    │◄───│  LOGISTICS   │◄───│ SALES ORDER  │
  │   (PDF Gen)  │    │  (shipped)   │    │(order_conf.) │
  │              │    │              │    │              │
  └──────────────┘    └──────────────┘    └──────────────┘
        ▲                    ▲                   │
        │                    │                   ▼
  ┌──────────────┐    ┌──────────────┐    ┌──────────────┐
  │   CLOSED     │    │   QUALITY    │◄───│ WORK ORDER   │
  │              │    │  (inspected) │    │(in_production)│
  │              │    │              │    │              │
  └──────────────┘    └──────────────┘    └──────────────┘
                                                │
                           Parallel:            │
                      ┌────────────────┐        │
                      │  VENDOR PO     │◄───────┘
                      │  (Procurement) │
                      │  RFQ → Quote   │
                      │  → PO → Receive│
                      └────────────────┘
''')

add_colored_heading('9.2 Procurement Flow', 2)
add_diagram_box('''
Procurement Workflow:
═══════════════════════════════════════════════════════════════

  Need Materials
       │
       ▼
  ┌──────────────┐    ┌──────────────┐    ┌──────────────┐
  │  CREATE RFQ  │───►│  SEND TO     │───►│  RECEIVE     │
  │  (Items +    │    │  VENDORS     │    │  VENDOR      │
  │   Specs)     │    │  (Multiple)  │    │  QUOTES      │
  └──────────────┘    └──────────────┘    └──────────────┘
                                                │
                                    Compare     │
                                    & Select    ▼
  ┌──────────────┐    ┌──────────────┐    ┌──────────────┐
  │  RECEIVE     │◄───│  GENERATE    │◄───│  APPROVE     │
  │  MATERIAL    │    │  PO          │    │  QUOTE       │
  │  (Stock +)   │    │  (PDF Gen)   │    │              │
  └──────────────┘    └──────────────┘    └──────────────┘
       │
       ▼
  Update MaterialStock
  + MaterialTransaction
''')

add_colored_heading('9.3 Multi-Tenant Data Flow', 2)
add_diagram_box('''
Request Lifecycle with Tenant Isolation:
═══════════════════════════════════════════════════════════════

  User (Company A)          Platform Admin
       │                         │
       │ GET /api/projects       │ GET /api/platform-admin/companies
       │ Auth: Bearer <token>    │ Auth: Bearer <token>
       │                         │
       ▼                         ▼
  ┌─────────────────────────────────────────────┐
  │              auth.js Middleware               │
  │  jwt.verify → req.user.company_id = "A"      │
  │  jwt.verify → req.user.role = "platform_admin"│
  └─────────────────────┬───────────────────────┘
                        │
                        ▼
  ┌─────────────────────────────────────────────┐
  │            tenant.js Middleware               │
  │                                              │
  │  Company A User:                             │
  │    WHERE company_id = 'A'  (auto-injected)   │
  │                                              │
  │  Platform Admin:                             │
  │    No company_id filter (sees all)           │
  └─────────────────────┬───────────────────────┘
                        │
                        ▼
  ┌─────────────────────────────────────────────┐
  │              Sequelize Query                  │
  │  Company A: SELECT * FROM projects            │
  │             WHERE company_id = 'A'            │
  │                                              │
  │  Platform:  SELECT * FROM companies           │
  │             (no tenant filter)               │
  └─────────────────────────────────────────────┘
''')

# ════════════════════════════════════════════════════════════════════════
#  10. DEPLOYMENT ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_colored_heading('10. Deployment Architecture', 1)

add_colored_heading('10.1 Infrastructure Diagram', 2)
add_diagram_box('''
Deployment Infrastructure:
═══════════════════════════════════════════════════════════════

  ┌──────────────────────────────────────────────────────┐
  │                    GITHUB                             │
  │         DevopsCDPL/forged-idas                       │
  │                                                      │
  │  ┌──────────┐   PR/Merge    ┌──────────┐            │
  │  │   dev    │──────────────►│   main   │            │
  │  │ (deploy) │               │(protected)│            │
  │  └────┬─────┘               └──────────┘            │
  └───────┼──────────────────────────────────────────────┘
          │
          │ Auto-deploy on push
          ▼
  ┌──────────────────────────────────────────────────────┐
  │                 RAILWAY CLOUD                         │
  │                                                      │
  │  ┌─────────────────┐    ┌─────────────────────┐     │
  │  │ FRONTEND SERVICE │    │  BACKEND SERVICE     │     │
  │  │                 │    │                     │     │
  │  │ Build:          │    │ Build:              │     │
  │  │  nixpacks       │    │  nixpacks           │     │
  │  │  CI=false npm   │    │  npm install        │     │
  │  │  run build      │    │                     │     │
  │  │                 │    │ Start:              │     │
  │  │ Start:          │    │  node src/index.js  │     │
  │  │  node start.js  │    │                     │     │
  │  │  (serves static │    │ Health:             │     │
  │  │   + API inject) │    │  GET /health        │     │
  │  │                 │    │                     │     │
  │  │ Health: GET /   │    │ Restart: on_failure │     │
  │  │ Restart: ×5     │    │ Max retries: 3      │     │
  │  └─────────────────┘    └──────────┬──────────┘     │
  │                                    │                 │
  │                                    │ Private Network │
  │                                    ▼                 │
  │                         ┌─────────────────────┐     │
  │                         │  POSTGRESQL          │     │
  │                         │  (Railway Managed)   │     │
  │                         │                     │     │
  │                         │  Connection:        │     │
  │                         │  Private domain     │     │
  │                         │  (internal)         │     │
  │                         │                     │     │
  │                         │  SSL: Required      │     │
  │                         │  Pool: 2-10 conn.   │     │
  │                         └─────────────────────┘     │
  └──────────────────────────────────────────────────────┘
''')

add_colored_heading('10.2 Frontend Runtime Injection', 2)
doc.add_paragraph('The frontend uses a unique runtime API URL injection mechanism via start.js:')

add_diagram_box('''
Frontend Serving (start.js):
═══════════════════════════════════════════════════════════════

  1. Read REACT_APP_API_URL or BACKEND_URL from env
  2. Inject into index.html as: window.__RUNTIME_API_URL__
  3. Serve static build/ via express on port $PORT
  4. Axios base URL reads window.__RUNTIME_API_URL__ at init

  This avoids hardcoding API URLs at build time,
  allowing the same build artifact to target different backends.
''')

add_colored_heading('10.3 Environment Variables', 2)
add_table(
    ['Variable', 'Service', 'Purpose'],
    [
        ['DATABASE_URL', 'Backend', 'PostgreSQL connection string'],
        ['JWT_SECRET', 'Backend', 'JWT signing key'],
        ['PORT', 'Both', 'Server port (Railway auto-sets)'],
        ['NODE_ENV', 'Both', 'production / development'],
        ['REACT_APP_API_URL', 'Frontend', 'Backend API base URL'],
        ['ADMIN_EMAIL', 'Backend', 'Self-healing admin account email'],
        ['ADMIN_PASSWORD', 'Backend', 'Self-healing admin account password'],
        ['SMTP_HOST/PORT/USER/PASS', 'Backend', 'Email service configuration'],
        ['CORS_ORIGINS', 'Backend', 'Additional CORS allowed origins'],
    ],
    [2.2, 1, 3.3]
)

# ════════════════════════════════════════════════════════════════════════
#  11. SECURITY ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_colored_heading('11. Security Architecture', 1)

add_diagram_box('''
Security Layers:
═══════════════════════════════════════════════════════════════

  ┌─────────────────────────────────────────────────┐
  │  Layer 1: TRANSPORT                              │
  │  • HTTPS/TLS (Railway managed)                   │
  │  • CORS whitelist                                │
  │  • Helmet security headers                       │
  └─────────────────────────────────────────────────┘
  ┌─────────────────────────────────────────────────┐
  │  Layer 2: RATE LIMITING                          │
  │  • Per-IP request throttling                     │
  │  • Account lock after failed logins              │
  └─────────────────────────────────────────────────┘
  ┌─────────────────────────────────────────────────┐
  │  Layer 3: AUTHENTICATION                         │
  │  • JWT tokens (signed with HS256)                │
  │  • bcryptjs password hashing (10 rounds)         │
  │  • OTP-based password reset                      │
  │  • Session tracking & monitoring                 │
  └─────────────────────────────────────────────────┘
  ┌─────────────────────────────────────────────────┐
  │  Layer 4: AUTHORIZATION                          │
  │  • Role-based access control (5 roles)           │
  │  • Hierarchical role system (100→10)             │
  │  • Co-Admin slot assignments (Owner/Co-Owner)    │
  │  • Module-level permissions (read/write/admin)   │
  │  • Custom roles (enterprise)                     │
  └─────────────────────────────────────────────────┘
  ┌─────────────────────────────────────────────────┐
  │  Layer 5: DATA ISOLATION                         │
  │  • Multi-tenant row-level filtering              │
  │  • Anti-spoofing (JWT company_id enforced)       │
  │  • Soft deletes (no permanent data loss)         │
  │  • Audit logging (all CRUD operations tracked)   │
  └─────────────────────────────────────────────────┘
  ┌─────────────────────────────────────────────────┐
  │  Layer 6: INPUT VALIDATION                       │
  │  • express-validator on all routes               │
  │  • File upload type/size restrictions             │
  │  • SQL injection protection (Sequelize ORM)      │
  │  • XSS prevention (React auto-escaping + Helmet) │
  └─────────────────────────────────────────────────┘
''')

# ════════════════════════════════════════════════════════════════════════
#  12. MODULE BREAKDOWN
# ════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_colored_heading('12. Module Breakdown', 1)
doc.add_paragraph('Forged-IDAS is organized into the following functional modules:')

add_table(
    ['Module', 'Frontend Pages', 'Backend Routes', 'Models', 'Key Features'],
    [
        ['Dashboard', '1', '2', '2', 'Stats, charts, recent activity'],
        ['Projects', '2', '1', '2', 'CRUD, status workflow, 12-tab detail view'],
        ['Estimation', '(tab)', '1', '1', 'Multi-revision estimates, approvals'],
        ['Quotation', '(tab)', '-', '-', 'PDF generation from estimates'],
        ['Sales Orders', '(tab)', '1', '1', 'Client PO management'],
        ['Work Orders', '(tab)', '1', '1', 'Production planning'],
        ['Quality', '(tab)', '1', '1', 'Inspection records, COC generation'],
        ['Logistics', '(tab)', '1', '1', 'Shipping and delivery tracking'],
        ['Invoicing', '(tab)', '-', '1', 'Invoice generation (PDF)'],
        ['Clients', '3', '1', '1', 'Client CRUD, contacts, history'],
        ['Vendors', '4', '1', '1', 'Vendor CRUD, materials, rating'],
        ['Parts Master', '1', '1', '3', 'Parts, dimensions, templates, heat tracking'],
        ['Raw Materials', '1', '1', '1', 'Material specifications, vendor mapping'],
        ['Material Stock', '1', '1', '3', 'Inventory tracking, transactions'],
        ['Procurement', '2', '3', '6', 'RFQ→Quote→PO workflow, bundles'],
        ['Vendor PO', '1', '2', '3', 'Vendor-specific purchase orders'],
        ['File Manager', '1', '1', '2', 'Project folders, document upload'],
        ['Access Control', '1 (8 tabs)', '3', '3', 'Users, roles, permissions, templates'],
        ['Analytics', '2', '2', '1', 'Business analytics, project metrics'],
        ['Enterprise', '4', '4', '4', 'Sessions, custom roles, approvals, risk'],
        ['Chat/Messages', '1', '1', '3', 'Real-time messaging'],
        ['Settings', '1', '1', '1', 'System configuration'],
        ['Recycle Bin', '1', '1', '-', 'Soft-delete recovery'],
        ['Platform Admin', '6+', '1', '-', 'Multi-tenant management, SaaS admin'],
    ],
    [1.2, 0.9, 1, 0.7, 2.5]
)

# ─── Footer Note ───────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Architecture Document —')
run.italic = True
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ─── Save ──────────────────────────────────────────────────────────────

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Forged-IDAS_Architecture.docx')
doc.save(output_path)
print(f'Document saved: {output_path}')
